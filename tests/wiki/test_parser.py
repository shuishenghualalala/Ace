import asyncio
import gc
import hashlib
import json
import os
import sys
import tempfile
import zipfile
from dataclasses import replace
from io import BytesIO
from pathlib import Path
from types import ModuleType, SimpleNamespace
from unittest.mock import MagicMock

import pytest
import psutil

from crew.security.context import SecurityContext
from crew.security.launch import (
    current_process_launch,
    finalize_process_launch,
    issue_process_launch,
    validate_process_launch,
)
from crew.security.models import (
    FilesystemAccess,
    FilesystemEntry,
    PermissionProfile,
    PermissionProfileKind,
    NetworkPolicy,
)
from crew.security.runtime_client import NativeRuntimeError, RuntimeErrorCode
from crew.tools.file_utils import FileConflictError
from crew.wiki.parser import (
    DocumentParseQualityError,
    MissingDependencyError,
    _markdown_table_from_rows,
    guess_mime_type,
    parse_document_from_bytes,
    parse_document_from_bytes_async,
    parse_document_to_markdown,
)


def _open_fds() -> int:
    gc.collect()
    process = psutil.Process()
    return process.num_handles() if os.name == "nt" else process.num_fds()


def test_parse_txt():
    with tempfile.TemporaryDirectory() as tmp:
        path = Path(tmp) / "test.txt"
        path.write_text("hello world", encoding="utf-8")
        assert parse_document_to_markdown(path) == "hello world"


def test_parse_md():
    with tempfile.TemporaryDirectory() as tmp:
        path = Path(tmp) / "test.md"
        path.write_text("# Title\n\nbody", encoding="utf-8")
        assert parse_document_to_markdown(path) == "# Title\n\nbody"


def test_parse_unknown_extension_falls_back_to_text():
    with tempfile.TemporaryDirectory() as tmp:
        path = Path(tmp) / "test.xyz"
        path.write_text("plain text", encoding="utf-8")
        assert parse_document_to_markdown(path) == "plain text"


def test_parse_missing_file_raises():
    with pytest.raises(FileNotFoundError):
        parse_document_to_markdown("/nonexistent/file.txt")


def test_parse_from_bytes():
    text = "from bytes"
    assert parse_document_from_bytes(text.encode("utf-8"), "memo.txt") == text


def test_parse_from_bytes_uses_pdf_signature_when_filename_has_no_extension(monkeypatch):
    import crew.wiki.parser as parser_mod

    monkeypatch.setattr(parser_mod, "_parse_pdf", lambda path: f"parsed:{path.suffix}")

    assert parse_document_from_bytes(b"%PDF-1.7\nfake", "display title") == "parsed:.pdf"


def test_parse_failure_cleans_up_every_temporary_directory(monkeypatch):
    import crew.wiki.parser as parser_mod

    created: list[Path] = []

    class TrackingTemporaryDirectory(tempfile.TemporaryDirectory):
        def __init__(self, *args, **kwargs):
            super().__init__(*args, **kwargs)
            created.append(Path(self.name))

    monkeypatch.setattr(tempfile, "TemporaryDirectory", TrackingTemporaryDirectory)

    def fail_pdf(_path):
        raise RuntimeError("simulated parser failure")

    monkeypatch.setattr(parser_mod, "_parse_pdf", fail_pdf)

    with pytest.raises(RuntimeError, match="simulated parser failure"):
        parse_document_from_bytes(b"%PDF-1.7\nfake", "display title")
    assert created
    for directory in created:
        assert not directory.exists()


def test_malformed_parse_paths_do_not_leak_file_descriptors():
    parse_document_from_bytes(b"warmup", "warmup.txt")
    baseline = _open_fds()
    for payload in (
        b"%PDF-1.7\n%truncated",
        b"PK\x03\x04" + b"\x00" * 200,
        b"PK\x05\x06" + b"\x00" * 40,
    ):
        try:
            parse_document_from_bytes(payload, "malformed")
        except Exception:
            pass
    assert _open_fds() <= baseline + 2


def test_parse_from_bytes_detects_ooxml_container_without_extension(monkeypatch):
    import crew.wiki.parser as parser_mod

    payload = BytesIO()
    with zipfile.ZipFile(payload, "w") as archive:
        archive.writestr("[Content_Types].xml", "<Types/>")
        archive.writestr("word/document.xml", "<document/>")
    monkeypatch.setattr(parser_mod, "_parse_docx", lambda path: f"parsed:{path.suffix}")

    assert parse_document_from_bytes(payload.getvalue(), "display title") == "parsed:.docx"


def test_parse_gb18030_text_detects_encoding():
    text = "中文知识库内容"
    assert parse_document_from_bytes(text.encode("gb18030"), "memo.txt") == text


def test_parse_unknown_binary_is_rejected():
    with pytest.raises(DocumentParseQualityError, match="二进制"):
        parse_document_from_bytes(b"\x00\x01\x02\x03\x00\x10", "data.bin")


def test_parse_replacement_character_quality_gate():
    with pytest.raises(DocumentParseQualityError, match="乱码比例"):
        parse_document_from_bytes(("正常文本" + "\ufffd" * 10).encode("utf-8"), "bad.txt")


def _converter_launch(tmp_path: Path, *, managed: bool = True):
    workspace = tmp_path / "workspace"
    workspace.mkdir(exist_ok=True)
    context = SecurityContext(
        os_user="host-user",
        owner_account_id="owner-a",
        workspace_id="workspace-a",
        workspace_root=workspace,
        session_id="session-a",
        request_id="request-a",
        task_id="task-a",
        cwd=workspace,
    )
    if not managed:
        return issue_process_launch(
            context,
            PermissionProfile(PermissionProfileKind.DISABLED),
        )
    runtime = tmp_path / "ace-security-runtime"
    runtime.write_bytes(b"test-runtime")
    runtime.with_name("runtime-manifest.json").write_text(
        json.dumps(
            {
                "schema": 2,
                "binary_name": runtime.name,
                "binary_sha256": hashlib.sha256(runtime.read_bytes()).hexdigest(),
            }
        ),
        encoding="utf-8",
    )
    return issue_process_launch(
        context,
        PermissionProfile(
            PermissionProfileKind.MANAGED,
            filesystem=(
                FilesystemEntry(
                    root=workspace,
                    access=FilesystemAccess.READ_WRITE,
                ),
            ),
            network=NetworkPolicy.UNRESTRICTED,
            allow_local_binding=True,
        ),
        helper_argv=(str(runtime),),
    )


def _install_fake_soffice(monkeypatch, tmp_path: Path) -> Path:
    import crew.wiki.parser as parser_mod

    install_dir = tmp_path / "libreoffice" / "program"
    install_dir.mkdir(parents=True)
    executable = install_dir / "soffice"
    executable.write_bytes(b"fake-soffice")
    executable.chmod(0o700)
    monkeypatch.setattr(
        parser_mod,
        "_discover_libreoffice",
        lambda _environment: executable,
    )
    return executable


def test_libreoffice_discovery_does_not_search_current_directory(
    monkeypatch,
    tmp_path,
):
    import crew.wiki.parser as parser_mod

    executable_name = "soffice.exe" if os.name == "nt" else "soffice"
    executable = tmp_path / executable_name
    executable.write_bytes(b"fake-soffice")
    executable.chmod(0o700)
    monkeypatch.chdir(tmp_path)

    assert (
        parser_mod._discover_libreoffice(
            {"PATH": os.pathsep.join(("", ".", "relative-bin"))}
        )
        is None
    )


def test_parse_legacy_xls_converts_with_libreoffice(monkeypatch, tmp_path):
    import crew.wiki.parser as parser_mod

    source = tmp_path / "legacy.xls"
    source.write_bytes(b"legacy excel")
    _install_fake_soffice(monkeypatch, tmp_path)
    monkeypatch.setattr(parser_mod, "_parse_xlsx", lambda _path: "## 工作表: Sheet1\n\n| A |\n| --- |\n| 1 |")

    async def fake_execute(args, **_kwargs):
        output_dir = Path(args[args.index("--outdir") + 1])
        source_path = Path(args[-1])
        (output_dir / f"{source_path.stem}.xlsx").write_bytes(b"converted")
        return SimpleNamespace(returncode=0, stdout="", stderr="")

    monkeypatch.setattr("crew.security.launch.execute_captured", fake_execute)
    token = current_process_launch.set(_converter_launch(tmp_path))
    try:
        assert "工作表" in parse_document_to_markdown(source)
    finally:
        current_process_launch.reset(token)


@pytest.mark.asyncio
async def test_legacy_converter_rejects_scratch_usage_over_budget(
    monkeypatch,
    tmp_path,
):
    import crew.wiki.parser as parser_mod

    source = tmp_path / "legacy.xls"
    source.write_bytes(b"legacy excel")
    _install_fake_soffice(monkeypatch, tmp_path)
    monkeypatch.setattr(parser_mod, "_MAX_CONVERTER_SCRATCH_BYTES", 24)

    async def fake_execute(args, **_kwargs):
        output_dir = Path(args[args.index("--outdir") + 1])
        (output_dir / "oversized-profile.dat").write_bytes(b"x" * 32)
        source_path = Path(args[-1])
        (output_dir / f"{source_path.stem}.xlsx").write_bytes(b"converted")
        return SimpleNamespace(returncode=0, stdout="", stderr="")

    monkeypatch.setattr("crew.security.launch.execute_captured", fake_execute)
    token = current_process_launch.set(_converter_launch(tmp_path))
    try:
        with pytest.raises(RuntimeError, match="临时目录超过"):
            await parse_document_from_bytes_async(b"legacy excel", "legacy.xls")
    finally:
        current_process_launch.reset(token)


@pytest.mark.asyncio
async def test_legacy_converter_uses_managed_boundary_without_ambient_credentials(
    monkeypatch,
    tmp_path,
):
    import crew.wiki.parser as parser_mod

    executable = _install_fake_soffice(monkeypatch, tmp_path)
    ambient_bin = tmp_path / "ambient-bin"
    ambient_bin.mkdir()
    monkeypatch.setenv("AWS_SECRET_ACCESS_KEY", "ambient-aws-secret")
    monkeypatch.setenv("VLM_API_KEY", "ambient-vlm-secret")
    monkeypatch.setenv("GITHUB_TOKEN", "ambient-github-secret")
    monkeypatch.setenv("COMSPEC", str(ambient_bin / "cmd.exe"))
    monkeypatch.setenv("PATH", str(ambient_bin))
    monkeypatch.setenv("SHELL", str(ambient_bin / "sh"))
    monkeypatch.setattr(
        parser_mod,
        "_parse_xlsx",
        lambda _path: "## 工作表: Sheet1\n\n| A |\n| --- |\n| 1 |",
    )
    captured = {}

    async def fake_execute(args, **kwargs):
        launch = current_process_launch.get()
        validate_process_launch(launch)
        assert launch is not None and launch.managed
        output_dir = Path(kwargs["cwd"])
        authorization = finalize_process_launch(
            launch,
            argv=tuple(args),
            cwd=output_dir,
            environment=kwargs["env_overrides"],
        )
        source_path = Path(args[-1])
        (output_dir / f"{source_path.stem}.xlsx").write_bytes(b"converted")
        captured.update(
            args=args,
            authorization=authorization,
            launch=launch,
            **kwargs,
        )
        return SimpleNamespace(returncode=0, stdout="", stderr="")

    monkeypatch.setattr("crew.security.launch.execute_captured", fake_execute)
    token = current_process_launch.set(_converter_launch(tmp_path, managed=True))
    try:
        result = await parse_document_from_bytes_async(b"legacy excel", "legacy.xls")
    finally:
        current_process_launch.reset(token)

    assert "工作表" in result
    assert captured["args"][0] == str(executable.resolve())
    assert captured["timeout"] == 120.0
    assert captured["max_output_bytes"] == 256 * 1024
    assert captured["env_overrides"] == captured["env"]
    assert "AWS_SECRET_ACCESS_KEY" not in captured["env"]
    assert "VLM_API_KEY" not in captured["env"]
    assert "GITHUB_TOKEN" not in captured["env"]
    assert "COMSPEC" not in captured["env"]
    assert "SHELL" not in captured["env"]
    converter_path = captured["env"]["PATH"].split(os.pathsep)
    assert converter_path[0] == str(executable.resolve().parent)
    assert str(ambient_bin) not in converter_path
    for name in (
        "APPDATA",
        "HOME",
        "LOCALAPPDATA",
        "TEMP",
        "TMP",
        "TMPDIR",
        "USERPROFILE",
    ):
        assert captured["env"][name] == str(captured["cwd"])
    scratch_entries = [
        entry
        for entry in captured["launch"].additional_permissions.filesystem
        if entry.root == captured["cwd"]
    ]
    assert len(scratch_entries) == 1
    assert scratch_entries[0].access is FilesystemAccess.READ_WRITE
    assert scratch_entries[0].escalatable is False
    assert captured["launch"].profile.filesystem == ()
    assert captured["launch"].profile.network is NetworkPolicy.RESTRICTED
    assert captured["launch"].profile.allow_local_binding is False
    assert captured["launch"].trusted_readable_roots == (
        executable.resolve().parent.parent,
    )
    assert captured["authorization"].snapshot.writable_roots == (
        str(captured["cwd"].resolve()),
    )
    assert captured["authorization"].snapshot.readable_roots == (
        str(executable.resolve().parent.parent),
    )


@pytest.mark.asyncio
async def test_legacy_converter_rejects_executable_in_writable_capability(
    monkeypatch,
    tmp_path,
):
    import crew.wiki.parser as parser_mod

    launch = _converter_launch(tmp_path, managed=True)
    workspace = launch.profile.filesystem[0].root
    executable = workspace / "soffice"
    executable.write_bytes(b"workspace-controlled-soffice")
    executable.chmod(0o700)
    monkeypatch.setattr(
        parser_mod,
        "_discover_libreoffice",
        lambda _environment: executable,
    )
    monkeypatch.setattr(
        "crew.security.launch.execute_captured",
        lambda *_args, **_kwargs: pytest.fail("writable converter must not start"),
    )

    token = current_process_launch.set(launch)
    try:
        with pytest.raises(RuntimeError, match="用户可写授权目录"):
            await parse_document_from_bytes_async(b"legacy excel", "legacy.xls")
    finally:
        current_process_launch.reset(token)


@pytest.mark.asyncio
async def test_legacy_converter_rejects_executable_changed_during_run(
    monkeypatch,
    tmp_path,
):
    executable = _install_fake_soffice(monkeypatch, tmp_path)

    async def fake_execute(_args, **_kwargs):
        executable.write_bytes(b"replaced-soffice")
        return SimpleNamespace(returncode=0, stdout="", stderr="")

    monkeypatch.setattr("crew.security.launch.execute_captured", fake_execute)
    token = current_process_launch.set(_converter_launch(tmp_path))
    try:
        with pytest.raises(RuntimeError, match="运行期间已变化"):
            await parse_document_from_bytes_async(b"legacy excel", "legacy.xls")
    finally:
        current_process_launch.reset(token)


@pytest.mark.asyncio
async def test_legacy_converter_refuses_missing_launch_context(monkeypatch):
    import crew.wiki.parser as parser_mod

    monkeypatch.setattr(
        parser_mod,
        "_discover_libreoffice",
        lambda *_args, **_kwargs: pytest.fail("converter discovery must not run"),
    )
    monkeypatch.setattr(
        "crew.security.launch.execute_captured",
        lambda *_args, **_kwargs: pytest.fail("converter must not start"),
    )
    token = current_process_launch.set(None)
    try:
        with pytest.raises(NativeRuntimeError) as caught:
            await parse_document_from_bytes_async(b"legacy excel", "legacy.xls")
    finally:
        current_process_launch.reset(token)
    assert caught.value.code is RuntimeErrorCode.SANDBOX_UNAVAILABLE


@pytest.mark.asyncio
async def test_legacy_converter_refuses_disabled_launch_boundary(
    monkeypatch,
    tmp_path,
):
    import crew.wiki.parser as parser_mod

    monkeypatch.setattr(
        parser_mod,
        "_discover_libreoffice",
        lambda *_args, **_kwargs: pytest.fail("converter discovery must not run"),
    )
    monkeypatch.setattr(
        "crew.security.launch.execute_captured",
        lambda *_args, **_kwargs: pytest.fail("converter must not start"),
    )
    token = current_process_launch.set(_converter_launch(tmp_path, managed=False))
    try:
        with pytest.raises(NativeRuntimeError) as caught:
            await parse_document_from_bytes_async(b"legacy excel", "legacy.xls")
    finally:
        current_process_launch.reset(token)
    assert caught.value.code is RuntimeErrorCode.SANDBOX_UNAVAILABLE


@pytest.mark.asyncio
async def test_legacy_converter_refuses_stale_launch_authority(monkeypatch, tmp_path):
    import crew.wiki.parser as parser_mod

    stale = replace(_converter_launch(tmp_path), authority_digest="0" * 64)
    monkeypatch.setattr(
        parser_mod,
        "_discover_libreoffice",
        lambda *_args, **_kwargs: pytest.fail("converter discovery must not run"),
    )
    monkeypatch.setattr(
        "crew.security.launch.execute_captured",
        lambda *_args, **_kwargs: pytest.fail("converter must not start"),
    )
    token = current_process_launch.set(stale)
    try:
        with pytest.raises(NativeRuntimeError) as caught:
            await parse_document_from_bytes_async(b"legacy excel", "legacy.xls")
    finally:
        current_process_launch.reset(token)
    assert caught.value.code is RuntimeErrorCode.SANDBOX_DENIED


@pytest.mark.asyncio
async def test_legacy_converter_rejects_symlink_output_escape(monkeypatch, tmp_path):
    import crew.wiki.parser as parser_mod

    _install_fake_soffice(monkeypatch, tmp_path)
    outside = tmp_path / "outside.xlsx"
    outside.write_bytes(b"outside-secret")

    async def fake_execute(args, **kwargs):
        output_dir = Path(kwargs["cwd"])
        converted = output_dir / f"{Path(args[-1]).stem}.xlsx"
        try:
            converted.symlink_to(outside)
        except OSError:
            pytest.skip("symlinks are unavailable on this platform")
        return SimpleNamespace(returncode=0, stdout="", stderr="")

    monkeypatch.setattr("crew.security.launch.execute_captured", fake_execute)
    monkeypatch.setattr(
        parser_mod,
        "_parse_xlsx",
        lambda _path: pytest.fail("escaped output must not be parsed"),
    )
    token = current_process_launch.set(_converter_launch(tmp_path))
    try:
        with pytest.raises(FileConflictError, match="符号链接|链接|reparse"):
            await parse_document_from_bytes_async(b"legacy excel", "legacy.xls")
    finally:
        current_process_launch.reset(token)
    assert outside.read_bytes() == b"outside-secret"


@pytest.mark.asyncio
async def test_legacy_converter_timeout_is_bounded(monkeypatch, tmp_path):
    _install_fake_soffice(monkeypatch, tmp_path)

    async def fake_execute(_args, **kwargs):
        assert kwargs["timeout"] == 120.0
        assert kwargs["max_output_bytes"] == 256 * 1024
        raise TimeoutError

    monkeypatch.setattr("crew.security.launch.execute_captured", fake_execute)
    token = current_process_launch.set(_converter_launch(tmp_path))
    try:
        with pytest.raises(RuntimeError, match="转换超时"):
            await parse_document_from_bytes_async(b"legacy excel", "legacy.xls")
    finally:
        current_process_launch.reset(token)


@pytest.mark.asyncio
async def test_legacy_converter_cancellation_reaches_execution_boundary(monkeypatch, tmp_path):
    _install_fake_soffice(monkeypatch, tmp_path)
    started = asyncio.Event()
    cancelled = asyncio.Event()

    async def fake_execute(_args, **kwargs):
        assert kwargs["timeout"] == 120.0
        assert kwargs["max_output_bytes"] == 256 * 1024
        started.set()
        try:
            await asyncio.Event().wait()
        except asyncio.CancelledError:
            cancelled.set()
            raise

    monkeypatch.setattr("crew.security.launch.execute_captured", fake_execute)
    token = current_process_launch.set(_converter_launch(tmp_path))
    try:
        task = asyncio.create_task(
            parse_document_from_bytes_async(b"legacy excel", "legacy.xls")
        )
        await asyncio.wait_for(started.wait(), timeout=1)
        task.cancel()
        with pytest.raises(asyncio.CancelledError):
            await task
    finally:
        current_process_launch.reset(token)
    assert cancelled.is_set()


@pytest.mark.asyncio
async def test_converter_cancellation_releases_launch_and_fds(
    monkeypatch,
    tmp_path,
):
    _install_fake_soffice(monkeypatch, tmp_path)
    started = asyncio.Event()
    cancelled = asyncio.Event()

    async def fake_execute(_args, **kwargs):
        del kwargs
        started.set()
        try:
            await asyncio.Event().wait()
        except asyncio.CancelledError:
            cancelled.set()
            raise

    monkeypatch.setattr("crew.security.launch.execute_captured", fake_execute)
    launch = _converter_launch(tmp_path)
    token = current_process_launch.set(launch)
    baseline = _open_fds()
    try:
        task = asyncio.create_task(
            parse_document_from_bytes_async(b"legacy excel", "legacy.xls")
        )
        await asyncio.wait_for(started.wait(), timeout=1)
        task.cancel()
        with pytest.raises(asyncio.CancelledError):
            await task
        assert current_process_launch.get() is launch
    finally:
        current_process_launch.reset(token)
    assert cancelled.is_set()
    assert _open_fds() <= baseline + 2


def test_guess_mime_type():
    assert guess_mime_type("x.pdf") == "application/pdf"
    assert guess_mime_type("x.md") == "text/markdown"
    assert guess_mime_type("x.docx") == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    assert guess_mime_type("x.doc") == "application/msword"
    assert guess_mime_type("x.xlsx") == "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
    assert guess_mime_type("x.xls") == "application/vnd.ms-excel"
    assert guess_mime_type("x.pptx") == "application/vnd.openxmlformats-officedocument.presentationml.presentation"
    assert guess_mime_type("x.ppt") == "application/vnd.ms-powerpoint"
    assert guess_mime_type("x.jpg") == "image/jpeg"
    assert guess_mime_type("x.jpeg") == "image/jpeg"
    assert guess_mime_type("x.png") == "image/png"
    assert guess_mime_type("x.webp") == "image/webp"
    assert guess_mime_type("x.mp4") == "video/mp4"
    assert guess_mime_type("x.mov") == "video/quicktime"
    assert guess_mime_type("x.unknown") == "application/octet-stream"
    assert guess_mime_type("display title", b"%PDF-1.7\nfake") == "application/pdf"
    assert guess_mime_type("wrong.txt", b"\x89PNG\r\n\x1a\nfake") == "image/png"
    assert guess_mime_type("wrong.txt", b"\x00\x00\x00\x18ftypisom") == "video/mp4"


def _make_fake_openpyxl():
    """构造一个可返回简单工作表数据的 fake openpyxl 模块。"""
    fake = ModuleType("openpyxl")

    def load_workbook(path, data_only=True, read_only=False):
        wb = MagicMock()
        sheet = MagicMock()
        sheet.title = "Sheet1"
        sheet.iter_rows.return_value = [
            ("Name", "Age"),
            ("Alice", 30),
            ("Bob", 25),
            (None, None),  # 空行应被过滤
        ]
        wb.worksheets = [sheet]
        return wb

    fake.load_workbook = load_workbook
    return fake


def _make_fake_pptx():
    """构造一个可返回简单幻灯片文本的 fake python-pptx 模块。"""
    fake = ModuleType("pptx")

    def Presentation(path):
        prs = MagicMock()
        slide = MagicMock()
        slide.slide_layout.name = "标题幻灯片"
        slide.has_notes_slide = False
        shape = MagicMock()
        shape.has_text_frame = True
        shape.has_table = False
        shape.has_chart = False
        p1, p2 = MagicMock(), MagicMock()
        p1.text = "  Slide title  "
        p2.text = "Bullet one"
        shape.text_frame.paragraphs = [p1, p2]
        slide.shapes = [shape]
        prs.slides = [slide]
        return prs

    fake.Presentation = Presentation
    return fake


def _make_fake_pptx_full():
    """构造一个包含布局、文本、表格、图表、图片、备注的 fake python-pptx 模块。"""
    fake = ModuleType("pptx")
    fake.__path__ = []  # 让 from import 认为它是包
    enum_module = ModuleType("pptx.enum")
    enum_module.__path__ = []
    shapes_enum = ModuleType("pptx.enum.shapes")
    shapes_enum.__path__ = []

    class MSO_SHAPE_TYPE:
        PICTURE = 13

    shapes_enum.MSO_SHAPE_TYPE = MSO_SHAPE_TYPE
    enum_module.shapes = shapes_enum
    fake.enum = enum_module

    # 注册子模块，支持 from pptx.enum.shapes import MSO_SHAPE_TYPE
    sys.modules.setdefault("pptx", fake)
    sys.modules.setdefault("pptx.enum", enum_module)
    sys.modules.setdefault("pptx.enum.shapes", shapes_enum)

    def Presentation(path):
        prs = MagicMock()
        slide = MagicMock()
        slide.slide_layout.name = "标题幻灯片"

        notes_slide = MagicMock()
        notes_slide.notes_text_frame.text = "演讲者备注内容"
        slide.has_notes_slide = True
        slide.notes_slide = notes_slide

        # 文本 shape
        text_shape = MagicMock()
        text_shape.has_text_frame = True
        text_shape.has_table = False
        text_shape.has_chart = False
        text_shape.shape_type = None
        p1, p2 = MagicMock(), MagicMock()
        p1.text = "  Slide title  "
        p2.text = "Bullet one"
        text_shape.text_frame.paragraphs = [p1, p2]

        # 表格 shape
        table_shape = MagicMock()
        table_shape.has_text_frame = False
        table_shape.has_table = True
        table_shape.has_chart = False
        table_shape.shape_type = None
        cell_a1, cell_b1 = MagicMock(), MagicMock()
        cell_a1.text = "Name"
        cell_b1.text = "Age"
        cell_a2, cell_b2 = MagicMock(), MagicMock()
        cell_a2.text = "Alice"
        cell_b2.text = "30"
        row1, row2 = MagicMock(), MagicMock()
        row1.cells = [cell_a1, cell_b1]
        row2.cells = [cell_a2, cell_b2]
        table = MagicMock()
        table.rows = [row1, row2]
        table_shape.table = table

        # 图表 shape
        chart_shape = MagicMock()
        chart_shape.has_text_frame = False
        chart_shape.has_table = False
        chart_shape.has_chart = True
        chart_shape.shape_type = None
        cat1, cat2 = MagicMock(), MagicMock()
        cat1.label = "Q1"
        cat2.label = "Q2"
        series = MagicMock()
        series.name = "销售额"
        series.categories = [cat1, cat2]
        series.values = [100, 200]
        chart = MagicMock()
        chart.chart_type = "COLUMN_CLUSTERED"
        chart.series = [series]
        chart.has_data_labels = False
        chart_shape.chart = chart

        # 图片 shape
        pic_shape = MagicMock()
        pic_shape.has_text_frame = False
        pic_shape.has_table = False
        pic_shape.has_chart = False
        pic_shape.shape_type = MSO_SHAPE_TYPE.PICTURE
        image = MagicMock()
        image.ext = "png"
        image.blob = b"fake-image-bytes"
        pic_shape.image = image

        slide.shapes = [text_shape, table_shape, chart_shape, pic_shape]
        prs.slides = [slide]
        return prs

    fake.Presentation = Presentation
    return fake


def _make_fake_pptx_with_two_images():
    """构造一个包含两张图片的 fake python-pptx 模块，用于验证 failure cache。"""
    fake = ModuleType("pptx")
    fake.__path__ = []
    enum_module = ModuleType("pptx.enum")
    enum_module.__path__ = []
    shapes_enum = ModuleType("pptx.enum.shapes")
    shapes_enum.__path__ = []

    class MSO_SHAPE_TYPE:
        PICTURE = 13

    shapes_enum.MSO_SHAPE_TYPE = MSO_SHAPE_TYPE
    enum_module.shapes = shapes_enum
    fake.enum = enum_module
    sys.modules.setdefault("pptx", fake)
    sys.modules.setdefault("pptx.enum", enum_module)
    sys.modules.setdefault("pptx.enum.shapes", shapes_enum)

    def Presentation(path):
        prs = MagicMock()
        slide = MagicMock()
        slide.slide_layout.name = "标题幻灯片"
        slide.has_notes_slide = False

        def _pic():
            pic = MagicMock()
            pic.has_text_frame = False
            pic.has_table = False
            pic.has_chart = False
            pic.shape_type = MSO_SHAPE_TYPE.PICTURE
            image = MagicMock()
            image.ext = "png"
            image.blob = b"fake-image-bytes"
            pic.image = image
            return pic

        slide.shapes = [_pic(), _pic()]
        prs.slides = [slide]
        return prs

    fake.Presentation = Presentation
    return fake


def test_parse_xlsx(monkeypatch):
    monkeypatch.setitem(sys.modules, "openpyxl", _make_fake_openpyxl())
    with tempfile.TemporaryDirectory() as tmp:
        path = Path(tmp) / "test.xlsx"
        path.write_bytes(b"fake xlsx bytes")
        result = parse_document_to_markdown(path)
    assert "## 工作表: Sheet1" in result
    assert "| Name | Age |" in result
    assert "| Alice | 30 |" in result
    assert "| Bob | 25 |" in result


def test_parse_pptx(monkeypatch):
    monkeypatch.setitem(sys.modules, "pptx", _make_fake_pptx())
    with tempfile.TemporaryDirectory() as tmp:
        path = Path(tmp) / "test.pptx"
        path.write_bytes(b"fake pptx bytes")
        result = parse_document_to_markdown(path)
    assert "## 第 1 页" in result
    assert "**布局**: 标题幻灯片" in result
    assert "### 文本" in result
    assert "Slide title" in result
    assert "Bullet one" in result


def test_parse_pptx_full_elements(monkeypatch):
    """验证 PPTX 解析能输出布局、文本、表格、图表、图片（VLM 描述+OCR）与备注。"""
    monkeypatch.setitem(sys.modules, "pptx", _make_fake_pptx_full())
    monkeypatch.setattr(
        "crew.wiki.multimodal.describe_image",
        lambda path, prompt=None: "图片描述：一张示例示意图；图中文字：示例文字",
    )

    with tempfile.TemporaryDirectory() as tmp:
        path = Path(tmp) / "test.pptx"
        path.write_bytes(b"fake pptx bytes")
        result = parse_document_to_markdown(path)

    assert "**布局**: 标题幻灯片" in result
    assert "### 文本" in result
    assert "Slide title" in result
    assert "### 表格" in result
    assert "| Name | Age |" in result
    assert "| Alice | 30 |" in result
    assert "### 图表" in result
    assert "销售额" in result
    assert "Q1" in result
    assert "100" in result
    assert "### 图片" in result
    assert "图片描述：一张示例示意图" in result
    assert "示例文字" in result
    assert "### 备注" in result
    assert "演讲者备注内容" in result


def test_parse_pptx_image_failure_shows_user_hint(monkeypatch):
    """VLM 未配置/不可用或调用失败时，解析结果应包含用户可见的提示。"""
    from crew.wiki.multimodal import MediaUnderstandingError

    monkeypatch.setitem(sys.modules, "pptx", _make_fake_pptx_full())
    monkeypatch.setattr(
        "crew.wiki.multimodal.describe_image",
        lambda path, prompt=None: (_ for _ in ()).throw(
            MediaUnderstandingError("未找到 VLM_API_KEY，请在 .env 中配置")
        ),
    )

    with tempfile.TemporaryDirectory() as tmp:
        path = Path(tmp) / "test.pptx"
        path.write_bytes(b"fake pptx bytes")
        result = parse_document_to_markdown(path)

    assert "### 图片" in result
    assert "[图片内容未解析：未找到 VLM_API_KEY，请在 .env 中配置]" in result


def test_parse_pptx_image_failure_redacts_host_path_and_token(monkeypatch):
    from crew.wiki.multimodal import MediaUnderstandingError

    monkeypatch.setitem(sys.modules, "pptx", _make_fake_pptx_full())
    monkeypatch.setattr(
        "crew.wiki.multimodal.describe_image",
        lambda path, prompt=None: (_ for _ in ()).throw(
            MediaUnderstandingError(r"C:\private\key.pem ACCESS_TOKEN=must-not-leak")
        ),
    )

    with tempfile.TemporaryDirectory() as tmp:
        path = Path(tmp) / "test.pptx"
        path.write_bytes(b"fake pptx bytes")
        result = parse_document_to_markdown(path)

    assert "[图片内容未解析" in result
    assert r"C:\private\key.pem" not in result
    assert "must-not-leak" not in result


def test_parse_pptx_image_failure_uses_cache(monkeypatch):
    """同一次 PPT 解析中，第一张图片失败后后续图片应复用提示，不再重复调用 VLM。"""
    from crew.wiki.multimodal import MediaUnderstandingError

    monkeypatch.setitem(sys.modules, "pptx", _make_fake_pptx_with_two_images())
    call_count = 0

    def _fake_describe(path, prompt=None):
        nonlocal call_count
        call_count += 1
        raise MediaUnderstandingError("VLM 服务不可用")

    monkeypatch.setattr("crew.wiki.multimodal.describe_image", _fake_describe)

    with tempfile.TemporaryDirectory() as tmp:
        path = Path(tmp) / "test.pptx"
        path.write_bytes(b"fake pptx bytes")
        result = parse_document_to_markdown(path)

    # 两张图片，但 describe_image 只应被调用一次
    assert call_count == 1
    assert result.count("[图片内容未解析：VLM 服务不可用]") == 2


def test_parse_xlsx_prefers_read_only(monkeypatch):
    """xlsx 应优先用 read_only 模式解析，避免物化海量空样式单元格卡死。"""
    fake = ModuleType("openpyxl")
    calls: list[dict] = []

    def load_workbook(path, data_only=True, read_only=False):
        calls.append({"data_only": data_only, "read_only": read_only})
        wb = MagicMock()
        sheet = MagicMock()
        sheet.title = "Sheet1"
        sheet.iter_rows.return_value = [("Name", "Age"), ("Alice", 30)]
        wb.worksheets = [sheet]
        return wb

    fake.load_workbook = load_workbook
    monkeypatch.setitem(sys.modules, "openpyxl", fake)

    with tempfile.TemporaryDirectory() as tmp:
        path = Path(tmp) / "test.xlsx"
        path.write_bytes(b"fake xlsx bytes")
        result = parse_document_to_markdown(path)

    assert len(calls) == 1
    assert calls[0]["read_only"] is True
    assert "| Name | Age |" in result


def test_parse_xlsx_falls_back_to_normal_mode_on_read_only_failure(monkeypatch):
    """read_only 解析失败且无可修复的已知问题时，应降级到普通模式。"""
    fake = ModuleType("openpyxl")
    calls: list[dict] = []

    def load_workbook(path, data_only=True, read_only=False):
        calls.append({"data_only": data_only, "read_only": read_only})
        if read_only:
            raise Exception("read_only 无法读取")
        wb = MagicMock()
        sheet = MagicMock()
        sheet.title = "Sheet1"
        sheet.iter_rows.return_value = [
            ("Name", "Age"),
            ("Alice", 30),
            ("Bob", 25),
        ]
        wb.worksheets = [sheet]
        return wb

    fake.load_workbook = load_workbook
    monkeypatch.setitem(sys.modules, "openpyxl", fake)

    with tempfile.TemporaryDirectory() as tmp:
        path = Path(tmp) / "test.xlsx"
        path.write_bytes(b"fake xlsx bytes")
        result = parse_document_to_markdown(path)

    assert len(calls) == 2
    assert calls[0]["read_only"] is True
    assert calls[1]["read_only"] is False
    assert "## 工作表: Sheet1" in result
    assert "| Name | Age |" in result


def test_parse_xlsx_truncates_oversized_sheet(monkeypatch):
    """超过行/列上限的工作表应被截断并标注，防止病态文件拖垮解析。"""
    import crew.wiki.parser as parser_mod

    monkeypatch.setattr(parser_mod, "_XLSX_MAX_ROWS_PER_SHEET", 5)
    monkeypatch.setattr(parser_mod, "_XLSX_MAX_COLS", 3)
    fake = ModuleType("openpyxl")

    def load_workbook(path, data_only=True, read_only=False):
        wb = MagicMock()
        sheet = MagicMock()
        sheet.title = "Big"
        # 10 行 x 5 列真实数据，行/列都超过调低后的上限
        sheet.iter_rows.return_value = [
            tuple(f"r{r}c{c}" for c in range(5)) for r in range(10)
        ]
        wb.worksheets = [sheet]
        return wb

    fake.load_workbook = load_workbook
    monkeypatch.setitem(sys.modules, "openpyxl", fake)

    with tempfile.TemporaryDirectory() as tmp:
        path = Path(tmp) / "big.xlsx"
        path.write_bytes(b"fake xlsx bytes")
        result = parse_document_to_markdown(path)

    assert "## 工作表: Big" in result
    # 只保留 5 行、3 列
    assert "| r0c0 | r0c1 | r0c2 |" in result
    assert "r0c3" not in result
    assert "r5c0" not in result
    assert "已截断" in result


@pytest.mark.parametrize("pages,objects", [(3, 1), (1, 3)])
def test_parse_pdf_rejects_page_or_object_bombs(monkeypatch, pages, objects):
    import crew.wiki.parser as parser_mod

    class FakeDocument:
        page_count = pages

        def xref_length(self):
            return objects

        def close(self):
            pass

    fitz = ModuleType("fitz")
    fitz.open = lambda _path: FakeDocument()
    monkeypatch.setitem(sys.modules, "fitz", fitz)
    monkeypatch.setattr(parser_mod, "_MAX_PDF_PAGES", 2)
    monkeypatch.setattr(parser_mod, "_MAX_PDF_OBJECTS", 2)

    with pytest.raises(DocumentParseQualityError, match="安全上限"):
        parser_mod._parse_pdf(Path("bomb.pdf"))


def test_parse_xlsx_ignores_dimension_padding_beyond_col_cap(monkeypatch):
    """dimension 把行补齐到整行宽度（全 None）时不应误报截断。"""
    import crew.wiki.parser as parser_mod

    monkeypatch.setattr(parser_mod, "_XLSX_MAX_COLS", 3)
    fake = ModuleType("openpyxl")

    def load_workbook(path, data_only=True, read_only=False):
        wb = MagicMock()
        sheet = MagicMock()
        sheet.title = "Sheet1"
        # read_only 在 dimension 过宽时会把每行补齐到该宽度（补 None）
        sheet.iter_rows.return_value = [
            ("Name", "Age", "City", None, None),
            ("Alice", 30, "BJ", None, None),
        ]
        wb.worksheets = [sheet]
        return wb

    fake.load_workbook = load_workbook
    monkeypatch.setitem(sys.modules, "openpyxl", fake)

    with tempfile.TemporaryDirectory() as tmp:
        path = Path(tmp) / "padded.xlsx"
        path.write_bytes(b"fake xlsx bytes")
        result = parse_document_to_markdown(path)

    assert "| Name | Age | City |" in result
    assert "已截断" not in result
    # 尾部补齐的空列应被裁掉，不产生空单元格
    assert "|  |" not in result


@pytest.mark.parametrize(
    "module_name,extension,dependency",
    [
        ("openpyxl", "xlsx", "openpyxl"),
        ("pptx", "pptx", "python-pptx"),
        ("docx", "docx", "python-docx"),
    ],
)
def test_parse_missing_dependency(monkeypatch, module_name, extension, dependency):
    """未安装对应依赖时应抛出 MissingDependencyError。"""
    monkeypatch.setitem(sys.modules, module_name, None)
    with tempfile.TemporaryDirectory() as tmp:
        path = Path(tmp) / f"test.{extension}"
        path.write_bytes(b"fake bytes")
        with pytest.raises(MissingDependencyError) as exc_info:
            parse_document_to_markdown(path)
    assert exc_info.value.dependency == dependency


def test_parse_xlsx_repairs_empty_fill_tags():
    """遇到 Excel/WPS 生成的空 <fill/> 标签时，_parse_xlsx 应在原有失败兜底链里自动修复并解析成功。"""
    import zipfile

    openpyxl = pytest.importorskip("openpyxl")

    from crew.wiki.parser import parse_document_to_markdown

    with tempfile.TemporaryDirectory() as tmp:
        valid = Path(tmp) / "valid.xlsx"
        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = "Sheet1"
        ws["A1"] = "Name"
        ws["A2"] = "Alice"
        wb.save(valid)

        # 往 styles.xml 注入空 <fill/>，构造 openpyxl 正常模式会失败的文件
        patched = Path(tmp) / "patched.xlsx"
        tmp_dir = Path(tmp) / "xlsx_tmp"
        with zipfile.ZipFile(valid, "r") as zf:
            zf.extractall(tmp_dir)
        styles = tmp_dir / "xl" / "styles.xml"
        text = styles.read_text(encoding="utf-8")
        text = text.replace("</fills>", '<fill/></fills>')
        styles.write_text(text, encoding="utf-8")
        with zipfile.ZipFile(patched, "w", zipfile.ZIP_DEFLATED) as zout:
            for f in tmp_dir.rglob("*"):
                if f.is_file():
                    zout.write(f, f.relative_to(tmp_dir))

        result = parse_document_to_markdown(patched)
        assert "## 工作表: Sheet1" in result
        assert "| Name |" in result
        assert "| Alice |" in result


def test_markdown_table_from_rows():
    """共享表格转 Markdown  helper 应输出标准表格并过滤空行。"""
    rows = [["指标", "数值"], ["A", "1"], ["B", "2"]]
    result = _markdown_table_from_rows(rows)
    assert result is not None
    assert "| 指标 | 数值 |" in result
    assert "| A | 1 |" in result
    assert "| B | 2 |" in result
    assert "| --- | --- |" in result


def test_markdown_table_from_rows_filters_empty_rows_and_pads_columns():
    """空行应被过滤，列数不一致时应补齐空列。"""
    rows = [["指标", "数值"], ["", ""], ["A", "1", "额外"]]
    result = _markdown_table_from_rows(rows)
    assert result is not None
    # 第三行被补齐到 3 列
    assert "| A | 1 | 额外 |" in result
    # 空行被过滤，结果只剩表头+一行数据
    assert result.count("\n") == 2


def _make_fake_docx_with_table():
    """构造一个包含段落和表格的 fake python-docx 模块，用于验证 DOCX 解析顺序。"""
    fake = ModuleType("docx")
    fake.__path__ = []

    fake_text = ModuleType("docx.text")
    fake_text.__path__ = []

    class FakeParagraph:
        def __init__(self, element, parent):
            self._text = getattr(element, "_text", "")

        @property
        def text(self):
            return self._text

    fake_paragraph = ModuleType("docx.text.paragraph")
    fake_paragraph.Paragraph = FakeParagraph
    fake_text.paragraph = fake_paragraph

    class FakeCell:
        def __init__(self, text):
            self.text = text

    class FakeRow:
        def __init__(self, cells):
            self.cells = [FakeCell(c) for c in cells]

    class FakeTable:
        def __init__(self, element, parent):
            self._rows = getattr(element, "_rows", [])

        @property
        def rows(self):
            return self._rows

    fake_table = ModuleType("docx.table")
    fake_table.Table = FakeTable

    # fake oxml.ns 供 _parse_docx import
    fake_oxml = ModuleType("docx.oxml")
    fake_oxml.__path__ = []
    fake_oxml_ns = ModuleType("docx.oxml.ns")
    fake_oxml_ns.__path__ = []

    def _qn(tag):
        local = tag.split(":")[-1]
        return f"{{http://schemas.openxmlformats.org/wordprocessingml/2006/main}}{local}"

    fake_oxml_ns.qn = _qn
    fake_oxml.ns = fake_oxml_ns

    class FakeParagraphElement:
        tag = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}p"

        def __init__(self, text):
            self._text = text

    class FakeTableElement:
        tag = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tbl"

        def __init__(self, rows):
            self._rows = [FakeRow(r) for r in rows]

    class FakeBody:
        def __init__(self, children):
            self._children = children

        def __iter__(self):
            return iter(self._children)

    class FakeElement:
        def __init__(self, body):
            self.body = body

    class FakeDocument:
        def __init__(self, path):
            self.element = FakeElement(
                FakeBody(
                    [
                        FakeParagraphElement("省经规范前言"),
                        FakeParagraphElement(""),  # 空段落应被忽略
                        FakeTableElement([["指标", "数值"], ["A", "1"]]),
                        FakeParagraphElement("省经规范结语"),
                    ]
                )
            )
            # _parse_docx 优先使用 document._body 作为 Paragraph/Table 的 parent
            self._body = self

    fake.Document = FakeDocument
    fake.text = fake_text
    fake.table = fake_table
    fake.oxml = fake_oxml

    return fake


def test_parse_docx_extracts_paragraphs_and_tables_in_order(monkeypatch):
    """DOCX 解析应按原文档顺序输出段落和表格，不能丢失表格。"""
    fake = _make_fake_docx_with_table()
    monkeypatch.setitem(sys.modules, "docx", fake)
    monkeypatch.setitem(sys.modules, "docx.text", fake.text)
    monkeypatch.setitem(sys.modules, "docx.text.paragraph", fake.text.paragraph)
    monkeypatch.setitem(sys.modules, "docx.table", fake.table)
    monkeypatch.setitem(sys.modules, "docx.oxml", fake.oxml)
    monkeypatch.setitem(sys.modules, "docx.oxml.ns", fake.oxml.ns)
    with tempfile.TemporaryDirectory() as tmp:
        path = Path(tmp) / "test.docx"
        path.write_bytes(b"fake docx bytes")
        result = parse_document_to_markdown(path)

    assert "省经规范前言" in result
    assert "省经规范结语" in result
    assert "| 指标 | 数值 |" in result
    assert "| A | 1 |" in result
    # 表格应夹在段落中间，而不是被丢弃
    assert result.index("省经规范前言") < result.index("| 指标 | 数值 |")
    assert result.index("| 指标 | 数值 |") < result.index("省经规范结语")


def test_parse_docx_extracts_table_inside_content_control():
    """DOCX 中的表格若被 <w:sdt> 内容控件包裹，解析时不能丢失。"""
    docx = pytest.importorskip("docx")
    from docx.oxml import OxmlElement

    with tempfile.TemporaryDirectory() as tmp:
        path = Path(tmp) / "sdt_table.docx"
        document = docx.Document()
        document.add_paragraph("普通段落")

        body = document.element.body
        sdt = OxmlElement("w:sdt")
        sdtPr = OxmlElement("w:sdtPr")
        sdtContent = OxmlElement("w:sdtContent")
        sdt.append(sdtPr)
        sdt.append(sdtContent)

        # 内容控件中的段落
        p = OxmlElement("w:p")
        r = OxmlElement("w:r")
        t = OxmlElement("w:t")
        t.text = "内容控件段落"
        r.append(t)
        p.append(r)
        sdtContent.append(p)

        # 内容控件中的表格（模拟字段明细表）
        tbl = OxmlElement("w:tbl")
        grid = OxmlElement("w:tblGrid")
        for _ in range(2):
            grid.append(OxmlElement("w:gridCol"))
        tbl.append(grid)
        for row_texts in [["字段", "类型"], ["USER_ID", "VARCHAR"]]:
            tr = OxmlElement("w:tr")
            for txt in row_texts:
                tc = OxmlElement("w:tc")
                tcPr = OxmlElement("w:tcPr")
                tc.append(tcPr)
                p_cell = OxmlElement("w:p")
                r_cell = OxmlElement("w:r")
                t_cell = OxmlElement("w:t")
                t_cell.text = txt
                r_cell.append(t_cell)
                p_cell.append(r_cell)
                tc.append(p_cell)
                tr.append(tc)
            tbl.append(tr)
        sdtContent.append(tbl)
        body.append(sdt)
        document.save(path)

        result = parse_document_to_markdown(path)

    assert "普通段落" in result
    assert "内容控件段落" in result
    assert "| 字段 | 类型 |" in result
    assert "| USER_ID | VARCHAR |" in result
    # 表格内容不应被重复提取为普通段落
    assert result.count("USER_ID") == 1
